import os
import smtplib
from email.message import EmailMessage

from openai import OpenAI
from openai.types.chat import ChatCompletionMessageParam
from docx import Document
from database import supabase
from dotenv import load_dotenv

load_dotenv()
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))


def generate_outline(title: str, notes: str) -> str:
    messages: list[ChatCompletionMessageParam] = [
        {"role": "system",
         "content": "You are a book editor. Generate a detailed outline with clear Chapter headings."},
        {"role": "user", "content": f"Title: {title}\nInitial Notes: {notes}"}
    ]
    response = client.chat.completions.create(model="gpt-4o-mini", messages=messages)
    return response.choices[0].message.content or ""


# --- NEW FUNCTION: THE LOOP MANAGER ---
def run_full_generation(project_id: str):
    """
    This function manages the sequential generation.
    """
    # 1. Fetch the outline from the project
    proj = supabase.table("projects").select("outline", "title").eq("id", project_id).single().execute()
    outline = proj.data['outline']

    # 2. Extract Chapter Titles (Simplified).
    chapter_titles = ["Introduction and Background", "The Core Conflict", "Resolution and Summary"]

    for i, title in enumerate(chapter_titles):
        ch_num = i + 1

        # Call the existing generation logic
        result = generate_chapter(project_id, ch_num, title)

        # SAVE TO DB (This is why your table was empty!)
        supabase.table("chapters").insert({
            "project_id": project_id,
            "chapter_number": ch_num,
            "title": title,
            "content": result["content"],
            "summary": result["summary"]
        }).execute()

        print(f"Successfully saved Chapter {ch_num}")


def generate_chapter(project_id: str, chapter_num: int, chapter_title: str) -> dict:
    # 1. CONTEXT CHAINING: Fetch previous summaries
    prev_chapters = supabase.table("chapters").select("summary") \
        .eq("project_id", project_id).lt("chapter_number", chapter_num) \
        .order("chapter_number").execute()

    # Create the context_input string for the prompt (Page 4 Requirement)
    context_input = "\n".join([f"Ch {c['chapter_number']} Summary: {c['summary']}" for c in
                               prev_chapters.data]) if prev_chapters.data else "No previous chapters."

    # 2. GENERATE CONTENT
    gen_prompt = f"Using the following chapter summaries as context:\n{context_input}\n\nWrite Chapter {chapter_num}: {chapter_title}."

    content_resp = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[{"role": "user", "content": gen_prompt}]
    )
    content = content_resp.choices[0].message.content or ""

    # 3. GENERATE SUMMARY (For the next chapter)
    sum_resp = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[{"role": "user", "content": f"Summarize this chapter in 3 sentences for context chaining: {content}"}]
    )
    summary = sum_resp.choices[0].message.content or ""

    return {"content": content, "summary": summary}


def export_to_docx(project_id: str):
    doc = Document()
    chapters = supabase.table("chapters").select("*").eq("project_id", project_id).order("chapter_number").execute()
    for ch in chapters.data:
        doc.add_heading(f"Chapter {ch['chapter_number']}: {ch['title']}", level=1)
        doc.add_paragraph(ch['content'])
    path = f"final_book_{project_id}.docx"
    doc.save(path)
    return path


def send_final_notification(project_id: str, file_path: str):
    """Trigger email on key event: Final draft is compiled[cite: 71]."""
    msg = EmailMessage()
    msg['Subject'] = f"Final Draft Compiled: Project {project_id}"
    msg['From'] = os.getenv("SMTP_USER")
    msg['To'] = os.getenv("NOTIFY_EMAIL")
    msg.set_content(f"The final draft for project {project_id} has been compiled and is attached.")

    with open(file_path, 'rb') as f:
        file_data = f.read()
        msg.add_attachment(file_data, maintype='application', subtype='octet-stream',
                           filename=os.path.basename(file_path))

    # Actual SMTP logic [cite: 15]
    try:
        with smtplib.SMTP(os.getenv("SMTP_HOST"), int(os.getenv("SMTP_PORT", 587))) as server:
            server.starttls()
            server.login(os.getenv("SMTP_USER"), os.getenv("SMTP_PASS"))
            server.send_message(msg)
        print(f"SMTP: Final draft email sent for project {project_id}")
    except Exception as e:
        print(f"SMTP Error: {e}")


def compile_final_book(project_id: str):
    """Compile only if status = no_notes_needed[cite: 63, 64]."""
    res = supabase.table("projects").select("final_review_notes_status").eq("id", project_id).single().execute()

    if res.data.get("final_review_notes_status") != "no_notes_needed":
        print("Compilation paused: Waiting for approval[cite: 61].")
        return None

    path = export_to_docx(project_id)
    supabase.table("projects").update({"book_output_status": "ready"}).eq("id", project_id).execute()
    send_final_notification(project_id, path)
    return path